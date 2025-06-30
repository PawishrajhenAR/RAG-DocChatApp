"""
Enhanced DOCX Reader

This module provides robust DOCX text extraction capabilities for RAG applications.
It handles various DOCX formats and ensures high-quality text extraction.
"""

import os
from typing import Dict, Any, List
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class EnhancedDOCXReader:
    """
    Enhanced DOCX reader with robust text extraction and metadata handling.
    
    Features:
    - Proper DOCX text extraction using python-docx
    - Metadata extraction from document properties
    - Error handling for corrupted or unsupported files
    - Text cleaning and formatting
    """
    
    def __init__(self):
        """Initialize the DOCX reader."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install with: pip install python-docx")
    
    def read_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Read and extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library not available. Install with: pip install python-docx")
        
        # Try primary extraction method
        try:
            return self._extract_with_python_docx(file_path)
        except Exception as primary_error:
            print(f"DEBUG: Primary DOCX extraction failed: {primary_error}")
            
            # Try fallback method
            try:
                print("DEBUG: Trying fallback DOCX extraction method...")
                return self._extract_with_fallback(file_path)
            except Exception as fallback_error:
                print(f"DEBUG: Fallback DOCX extraction also failed: {fallback_error}")
                raise Exception(f"All DOCX extraction methods failed. Primary error: {primary_error}, Fallback error: {fallback_error}")
    
    def _extract_with_python_docx(self, file_path: str) -> Dict[str, Any]:
        """Primary DOCX extraction method using python-docx."""
        print(f"DEBUG: Opening DOCX file: {file_path}")
        
        # Open the DOCX document
        doc = Document(file_path)
        
        print(f"DEBUG: DOCX opened successfully. Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
        
        # Extract text from paragraphs
        paragraphs = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
                if i < 3:  # Debug first few paragraphs
                    print(f"DEBUG: Paragraph {i}: {paragraph.text[:100]}...")
        
        # Extract text from tables
        tables = []
        for table_idx, table in enumerate(doc.tables):
            print(f"DEBUG: Processing table {table_idx + 1} with {len(table.rows)} rows")
            for row_idx, row in enumerate(table.rows):
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_row = ' | '.join(row_text)
                    tables.append(table_row)
                    if row_idx < 2:  # Debug first few rows
                        print(f"DEBUG: Table {table_idx + 1}, Row {row_idx + 1}: {table_row[:100]}...")
        
        # Combine all text
        all_text = []
        if paragraphs:
            all_text.extend(paragraphs)
        if tables:
            all_text.extend(tables)
        
        text = '\n\n'.join(all_text)
        
        # Clean the text
        text = self._clean_text(text)
        
        # Extract metadata
        metadata = self._extract_metadata(doc, file_path)
        
        result = {
            'text': text,
            'metadata': metadata,
            'paragraphs': len(paragraphs),
            'tables': len(doc.tables),
            'word_count': len(text.split()) if text else 0
        }
        
        print(f"DEBUG: DOCX extraction complete. Text length: {len(text)}, Words: {result['word_count']}")
        
        return result
    
    def _extract_with_fallback(self, file_path: str) -> Dict[str, Any]:
        """Fallback DOCX extraction method for corrupted files."""
        print("DEBUG: Using fallback DOCX extraction method")
        
        try:
            # Try to extract as ZIP file (DOCX is essentially a ZIP)
            import zipfile
            import xml.etree.ElementTree as ET
            
            text_parts = []
            
            with zipfile.ZipFile(file_path, 'r') as zip_file:
                # Look for the main document content
                if 'word/document.xml' in zip_file.namelist():
                    with zip_file.open('word/document.xml') as doc_xml:
                        tree = ET.parse(doc_xml)
                        root = tree.getroot()
                        
                        # Extract text from all text elements
                        for elem in root.iter():
                            if elem.text and elem.text.strip():
                                text_parts.append(elem.text.strip())
            
            text = '\n'.join(text_parts)
            text = self._clean_text(text)
            
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'file_type': 'docx',
                'extraction_method': 'fallback-zip'
            }
            
            result = {
                'text': text,
                'metadata': metadata,
                'paragraphs': 0,  # Unknown in fallback
                'tables': 0,      # Unknown in fallback
                'word_count': len(text.split()) if text else 0
            }
            
            print(f"DEBUG: Fallback DOCX extraction complete. Text length: {len(text)}, Words: {result['word_count']}")
            
            return result
            
        except Exception as e:
            raise Exception(f"Fallback DOCX extraction failed: {str(e)}")
    
    def _extract_metadata(self, doc: Document, file_path: str) -> Dict[str, Any]:
        """
        Extract metadata from the DOCX document.
        
        Args:
            doc: The DOCX document object
            file_path: Path to the file
            
        Returns:
            Dictionary containing document metadata
        """
        metadata = {
            'file_path': file_path,
            'file_name': os.path.basename(file_path),
            'file_type': 'docx',
            'extraction_method': 'python-docx'
        }
        
        try:
            # Extract core properties if available
            core_props = doc.core_properties
            if core_props.title:
                metadata['title'] = core_props.title
            if core_props.author:
                metadata['author'] = core_props.author
            if core_props.subject:
                metadata['subject'] = core_props.subject
            if core_props.created:
                metadata['created_date'] = core_props.created.isoformat()
            if core_props.modified:
                metadata['modified_date'] = core_props.modified.isoformat()
            
            # Count document elements
            metadata['paragraph_count'] = len(doc.paragraphs)
            metadata['table_count'] = len(doc.tables)
            metadata['section_count'] = len(doc.sections)
            
        except Exception as e:
            # If metadata extraction fails, continue with basic info
            print(f"Warning: Could not extract full metadata: {e}")
        
        return metadata
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip() 